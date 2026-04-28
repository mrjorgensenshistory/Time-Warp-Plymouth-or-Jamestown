"""
Build the student worksheet for Time Warp - Plymouth or Jamestown.
Output: Student-Worksheet-New-World.docx in the project root.
ARE-paragraph (Argument, Reasoning, Evidence) format, 8th grade.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT = Path(r"F:\Michael's\ACE\US History\Time-Warp-Plymouth-or-Jamestown")
OUT = PROJECT / "Student-Worksheet-New-World.docx"

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

# Default style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x2B, 0x4A)
    return p


def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p


def add_heading(text, size=14, color=(0x1A, 0x2B, 0x4A)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = True
    run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_para(text, size=11, italic=False, bold=False, indent=0.0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    run.bold = bold
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    return p


def add_blank_line(underline_chars=70):
    """Add an empty line of underscores for student to write on."""
    p = doc.add_paragraph()
    run = p.add_run("_" * underline_chars)
    run.font.size = Pt(11)
    return p


def add_box_field(label, blank_chars=50):
    """Inline label + underline."""
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.size = Pt(11)
    run2 = p.add_run("_" * blank_chars)
    run2.font.size = Pt(11)
    return p


# ===== HEADER =====

add_title("New World Comparison")
add_subtitle("An ARE Paragraph Essay  ·  Time Warp: Plymouth or Jamestown")

# Name / Date / Period row
header_p = doc.add_paragraph()
header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
header_p.add_run("Name: ").bold = True
header_p.add_run("_" * 30)
header_p.add_run("    Date: ").bold = True
header_p.add_run("_" * 14)
header_p.add_run("    Period: ").bold = True
header_p.add_run("_" * 6)

# Spacer
doc.add_paragraph()

# ===== INSTRUCTIONS =====

add_heading("Your Mission")
p = doc.add_paragraph()
p.add_run(
    "You just played Time Warp: Plymouth or Jamestown. "
    "You crossed the Atlantic in someone else's body. You made decisions. Some of you died."
).font.size = Pt(11)

p2 = doc.add_paragraph()
p2.add_run(
    "Now use what you lived through to write a three-paragraph essay explaining "
    "three real differences between the Plymouth colony and the Jamestown colony. "
    "Each paragraph follows the ARE format you've been practicing all year."
).font.size = Pt(11)

# ===== ARE FORMAT REMINDER =====

add_heading("ARE Paragraph Format (Reminder)", size=12, color=(0x55, 0x55, 0x55))

are_table = doc.add_table(rows=3, cols=2)
are_table.style = "Light Grid Accent 1"
are_data = [
    ("A — Argument", "One clear sentence stating your point."),
    ("R — Reasoning", "Explain WHY that point matters. Why does this difference between Plymouth and Jamestown matter for understanding America?"),
    ("E — Evidence", "Use SPECIFIC moments from the game. Name characters. Name decisions. Name what happened. (\"Smith took corn from the Powhatan at gunpoint.\" \"Squanto came back from Spain to find his village empty.\")"),
]
for i, (col1, col2) in enumerate(are_data):
    cell1 = are_table.cell(i, 0)
    cell1.text = col1
    for run in cell1.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(11)
    cell2 = are_table.cell(i, 1)
    cell2.text = col2
    for run in cell2.paragraphs[0].runs:
        run.font.size = Pt(11)

doc.add_paragraph()

# ===== THESIS STATEMENT =====

add_heading("Your Thesis Statement")
p = doc.add_paragraph()
p.add_run(
    "Fill in the three differences you will write about. "
    "Pick three. There are more than three to choose from — that's the point."
).italic = True

doc.add_paragraph()

# Thesis with blanks
thesis = doc.add_paragraph()
thesis.alignment = WD_ALIGN_PARAGRAPH.LEFT
r1 = thesis.add_run('"The Plymouth and Jamestown colonies were different due to ')
r1.font.size = Pt(12)
thesis.add_run("_" * 28).font.size = Pt(12)
thesis.add_run(", ").font.size = Pt(12)
thesis.add_run("_" * 28).font.size = Pt(12)
thesis.add_run(", and ").font.size = Pt(12)
thesis.add_run("_" * 28).font.size = Pt(12)
r2 = thesis.add_run('."')
r2.font.size = Pt(12)

# Difference bank for student reference
doc.add_paragraph()
add_heading("Difference Bank (pick three)", size=12, color=(0x55, 0x55, 0x55))

bank_items = [
    "WHY THEY CAME — Jamestown came for gold, glory, and Virginia Company profits. Plymouth came for religious freedom from the Church of England.",
    "WHO CAME — Jamestown was single young men, soldiers of fortune, no families. Plymouth was families, women, and children.",
    "GEOGRAPHY → ECONOMY — Jamestown had warm fertile land, deep navigable rivers, and a long growing season → cash crops (tobacco) → mass low-skilled labor demand → indentured servants → enslaved Africans. Plymouth had rocky soil, short growing season, small inlets → subsistence farming, fishing, shipbuilding → no economic engine for mass slavery.",
    "RELATIONSHIP WITH NATIVES — Jamestown started hostile (took Powhatan land and food at gunpoint). Plymouth started cooperative (Squanto and Massasoit alliance, first Thanksgiving), but the Wampanoag eventually became dependent and lost sovereignty too.",
    "DISEASE & FAMINE — Jamestown's Starving Time (winter 1609–10) killed about 80% of the colony — cannibalism documented. Plymouth lost about half the Mayflower passengers in the first winter. Native populations were devastated by European disease in BOTH places.",
    "GOVERNMENT — Plymouth signed the Mayflower Compact, the first written self-government on what became US soil. Jamestown was a private company colony for years before the House of Burgesses (1619).",
]

for item in bank_items:
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run(item).font.size = Pt(10)

doc.add_page_break()

# ===== PARAGRAPH 1 =====

add_heading("Paragraph 1: First Difference")
add_box_field("ARGUMENT (one sentence — what is the first difference?)", 70)
add_blank_line()
add_box_field("REASONING (one to two sentences — why does this difference matter?)", 60)
add_blank_line()
add_blank_line()
add_box_field("EVIDENCE (one to two sentences — name a character, decision, or moment from the game)", 50)
add_blank_line()
add_blank_line()
doc.add_paragraph()

# ===== PARAGRAPH 2 =====

add_heading("Paragraph 2: Second Difference")
add_box_field("ARGUMENT", 70)
add_blank_line()
add_box_field("REASONING", 60)
add_blank_line()
add_blank_line()
add_box_field("EVIDENCE", 50)
add_blank_line()
add_blank_line()
doc.add_paragraph()

# ===== PARAGRAPH 3 =====

add_heading("Paragraph 3: Third Difference")
add_box_field("ARGUMENT", 70)
add_blank_line()
add_box_field("REASONING", 60)
add_blank_line()
add_blank_line()
add_box_field("EVIDENCE", 50)
add_blank_line()
add_blank_line()
doc.add_paragraph()

# ===== TURN-IN CHECKLIST =====

add_heading("Before You Turn It In — Check Yourself")
checklist = [
    "My thesis names three different things.",
    "Each paragraph has all three: Argument, Reasoning, Evidence.",
    "I named at least one character from the game in each Evidence sentence.",
    "I did NOT say things like \"the colonists were bad\" or \"the Natives were good.\" Both sides were people doing survival math.",
    "I did NOT say \"imagine if\" or \"the lesson is.\" I stated what happened.",
    "My handwriting is readable.",
]
for item in checklist:
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run(item).font.size = Pt(11)

# ===== FOOTER =====

doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run("Mr. Jorgensen  ·  8th Grade US History  ·  ACE Charter")
fr.font.size = Pt(9)
fr.italic = True
fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.save(str(OUT))
print(f"[OK] Worksheet written: {OUT}")
print(f"     Size: {OUT.stat().st_size:,} bytes")
